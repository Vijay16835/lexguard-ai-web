import os
import sys
import stat
import shutil
import urllib.request
import logging
import traceback
import subprocess
from typing import Optional

# Setup module-level logger
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Supabase singleton — created once at module load, reused on every request.
# ---------------------------------------------------------------------------
_supabase_client = None

def get_supabase():
    """Return the shared Supabase client, initialising it on first call."""
    global _supabase_client
    if _supabase_client is None:
        from supabase import create_client
        from app.core.config import settings
        _supabase_client = create_client(settings.SUPABASE_URL, settings.SUPABASE_KEY)
    return _supabase_client


class TextExtractionError(Exception):
    """Custom exception raised when text extraction fails."""
    pass


# ---------------------------------------------------------------------------
# File Type and Size Constraints
# ---------------------------------------------------------------------------
SUPPORTED_EXTENSIONS = {'pdf', 'doc', 'docx', 'txt', 'jpg', 'jpeg', 'png', 'bmp', 'tiff', 'webp'}
MAX_FILE_SIZE_MB = 20


def get_file_extension(filename: str) -> str:
    """Get clean file extension from filename."""
    _, ext = os.path.splitext(filename)
    return ext.lstrip('.').lower()


def validate_file(filename: str, file_size: int) -> tuple:
    """Validate file type and size. Returns (is_valid, error_message)."""
    ext = get_file_extension(filename)
    if ext not in SUPPORTED_EXTENSIONS:
        return False, f"Unsupported file type: .{ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}"
    
    size_mb = file_size / (1024 * 1024)
    if size_mb > MAX_FILE_SIZE_MB:
        return False, f"File too large: {size_mb:.1f}MB. Maximum: {MAX_FILE_SIZE_MB}MB"
    
    return True, ""


# ---------------------------------------------------------------------------
# File Parsers (PDF, DOC, DOCX, TXT)
# ---------------------------------------------------------------------------
def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF files using pdfplumber with PyPDF2 fallback."""
    try:
        import pdfplumber
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
        return text.strip()
    except Exception as e:
        logger.warning(f"pdfplumber failed to parse PDF: {e}. Falling back to PyPDF2.")
        try:
            import PyPDF2
            text = ""
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            return text.strip()
        except Exception as e2:
            logger.error(f"PyPDF2 fallback also failed: {e2}")
            raise TextExtractionError("PDF text extraction failed") from e2


def extract_text_from_doc(file_path: str) -> str:
    """Extract text from legacy DOC files using legacy-doc."""
    try:
        from legacy_doc import extract_text as legacy_extract
        with open(file_path, "rb") as f:
            result = legacy_extract(f.read())
            if hasattr(result, 'text'):
                text = result.text
            else:
                text = str(result)
        if not text.strip():
            raise TextExtractionError("Unable to extract readable text from document.")
        return text.strip()
    except TextExtractionError:
        raise
    except Exception as e:
        logger.error(f"DOC extraction error: {e}")
        raise TextExtractionError("Unsupported file structure") from e


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX files."""
    try:
        import docx
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract cell text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text += row_text + "\n"
        
        if not text.strip():
            raise TextExtractionError("Unable to extract readable text from document.")
            
        return text.strip()
    except TextExtractionError:
        raise
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise TextExtractionError("Unsupported file structure") from e


def extract_text_from_txt(file_path: str) -> str:
    """Extract text from plain text files."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read().strip()
        if not text:
            raise TextExtractionError("Unable to extract readable text from document.")
        return text
    except TextExtractionError:
        raise
    except Exception as e:
        logger.error(f"TXT extraction error: {e}")
        raise TextExtractionError("Unsupported file structure") from e


# ---------------------------------------------------------------------------
# Tesseract Bootstrapping and Path Resolution
# ---------------------------------------------------------------------------
def _bootstrap_static_tesseract() -> str:
    """
    Download and configure static Tesseract binary on Linux native environments if needed.
    Returns path to verified binary, or empty string on failure.
    """
    if not sys.platform.startswith("linux"):
        return ""
        
    from app.core.config import settings
    
    # Store binary and data files relative to settings.UPLOAD_DIR to keep it self-contained
    bin_dir = os.path.abspath(os.path.join(settings.UPLOAD_DIR, "..", "bin"))
    tess_exe = os.path.join(bin_dir, "tesseract-static")
    tessdata_dir = os.path.join(bin_dir, "tessdata")
    tessdata_file = os.path.join(tessdata_dir, "eng.traineddata")
    fallback_tessdata_file = os.path.join(bin_dir, "eng.traineddata")
    
    # Helper to download with timeout and user-agent spoofing
    def _download_file(url: str, dest_path: str):
        try:
            req = urllib.request.Request(
                url,
                headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
            )
            with urllib.request.urlopen(req, timeout=30) as response:
                os.makedirs(os.path.dirname(dest_path), exist_ok=True)
                with open(dest_path, 'wb') as out_file:
                    shutil.copyfileobj(response, out_file)
        except Exception as e:
            logger.error(f"[TESS-INSTALL] Failed to download {url} to {dest_path}: {e}")
            raise RuntimeError(f"Download failed: {e}") from e

    # 1. Download static binary if missing
    if not os.path.exists(tess_exe):
        logger.info("[TESS-INSTALL] Downloading static Tesseract binary from GitHub...")
        url_bin = "https://github.com/DanielMYT/tesseract-static/releases/download/tesseract-5.5.3/tesseract.x86_64"
        try:
            _download_file(url_bin, tess_exe)
            st = os.stat(tess_exe)
            os.chmod(tess_exe, st.st_mode | stat.S_IEXEC)
            logger.info(f"[TESS-INSTALL] Static binary downloaded and permissions set at {tess_exe}")
        except Exception as e:
            logger.error(f"[TESS-INSTALL] Error preparing static binary: {e}")
            return ""

    # 2. Download fast English language pack if missing
    if not os.path.exists(tessdata_file):
        logger.info("[TESS-INSTALL] Downloading eng.traineddata language pack...")
        url_lang = "https://github.com/tesseract-ocr/tessdata_fast/raw/main/eng.traineddata"
        try:
            _download_file(url_lang, tessdata_file)
            logger.info(f"[TESS-INSTALL] Language pack downloaded successfully at {tessdata_file}")
        except Exception as e:
            logger.error(f"[TESS-INSTALL] Error preparing language pack: {e}")
            return ""

    # 3. Direct parent directory fallback copy for Musl path resolving
    if os.path.exists(tessdata_file) and not os.path.exists(fallback_tessdata_file):
        try:
            shutil.copy2(tessdata_file, fallback_tessdata_file)
            logger.info(f"[TESS-INSTALL] Copied eng.traineddata to parent fallback path {fallback_tessdata_file}")
        except Exception as e:
            logger.warning(f"[TESS-INSTALL] Fallback copy failed: {e}")

    # Verify everything exists and return path
    if os.path.exists(tess_exe) and os.path.exists(tessdata_file):
        os.environ["TESSDATA_PREFIX"] = tessdata_dir
        return tess_exe
        
    return ""


def get_tesseract_path() -> str:
    """
    Resolves the active tesseract binary path. 
    Checks settings, system PATH, typical Windows paths, and boots static fallback on Render if needed.
    """
    from app.core.config import settings

    # 1. Check configured path first
    cfg_path = settings.TESSERACT_CMD
    if cfg_path and os.path.exists(cfg_path):
        return cfg_path
        
    # 2. Check system PATH
    sys_path = shutil.which("tesseract")
    if sys_path:
        return sys_path
        
    # 3. Check typical Windows installation paths
    if sys.platform.startswith("win"):
        win_paths = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"
        ]
        for wp in win_paths:
            if os.path.exists(wp):
                return wp

    # 4. Dynamic static binary fallback (Render native platform)
    fallback_path = _bootstrap_static_tesseract()
    if fallback_path:
        return fallback_path
        
    return "tesseract"


# ---------------------------------------------------------------------------
# OCR and Image Parsing
# ---------------------------------------------------------------------------
def clean_extracted_text(text: str) -> str:
    """Clean and normalize extracted text, removing control characters and excessive whitespaces."""
    if not text:
        return ""
    # Filter out control characters except tabs/newlines
    cleaned = "".join(ch for ch in text if ch.isprintable() or ch in ('\n', '\r', '\t'))
    return "\n".join(line.strip() for line in cleaned.splitlines() if line.strip()).strip()


def validate_image_file(file_path: str):
    """Validate that the file exists, is not empty, and is a valid image within bounds."""
    if not os.path.exists(file_path):
        raise TextExtractionError("File not found on server.")
    if os.path.getsize(file_path) == 0:
        raise TextExtractionError("Image file is empty.")
    
    from PIL import Image
    try:
        # Verify format and integrity without loading data
        with Image.open(file_path) as img:
            img.verify()
    except Exception as e:
        raise TextExtractionError(f"Invalid image format: {e}")


def preprocess_image_pillow(file_path: str):
    """
    Load image, transpose EXIF orientation, handle transparency, and convert to grayscale.
    """
    from PIL import Image, ImageOps, ImageFilter
    
    try:
        # Increase pixel limit for decompression protection (e.g. max 100MP)
        Image.MAX_IMAGE_PIXELS = 100_000_000
        
        with Image.open(file_path) as img:
            width, height = img.size
            if width * height > Image.MAX_IMAGE_PIXELS:
                raise TextExtractionError(f"Image resolution too large ({width}x{height}). Max: 100MP.")
            
            # Correct EXIF rotation
            img = ImageOps.exif_transpose(img)
            
            # Flatten alpha/transparency channels on white background
            if img.mode in ("RGBA", "LA") or (img.mode == "P" and "transparency" in img.info):
                bg = Image.new("RGBA", img.size, (255, 255, 255, 255))
                if img.mode == "P":
                    img = img.convert("RGBA")
                bg.paste(img, (0, 0), img)
                img = bg.convert("RGB")
            else:
                img = img.convert("RGB")
            
            # Grayscale & subtle sharpen
            gray = img.convert("L")
            sharpened = gray.filter(ImageFilter.SHARPEN)
            return sharpened.copy()
    except TextExtractionError:
        raise
    except Exception as e:
        raise TextExtractionError(f"Image preprocessing failed: {e}") from e


def extract_text_from_image(file_path: str) -> str:
    """Validate, preprocess, and run high-performance OCR on the specified image file."""
    from app.services.ocr_service import ocr_service
    return ocr_service.extract_text_from_image(file_path)


def ocr_pdf(file_path: str) -> str:
    """Perform OCR on a PDF by rendering pages as images and running OCR on them."""
    import pypdfium2 as pdfium
    import tempfile
    
    logger.info(f"Starting PDF OCR fallback for: {file_path}")
    text = ""
    pdf = None
    try:
        pdf = pdfium.PdfDocument(file_path)
        for i in range(len(pdf)):
            page = pdf.get_page(i)
            pil_image = page.render(scale=2).to_pil()
            
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                temp_img_path = tmp.name
            try:
                pil_image.save(temp_img_path)
                page_text = extract_text_from_image(temp_img_path)
                if page_text:
                    text += page_text + "\n\n"
            finally:
                if os.path.exists(temp_img_path):
                    try:
                        os.remove(temp_img_path)
                    except Exception as e:
                        logger.warning(f"Failed to delete temp PDF page image: {e}")
    except Exception as e:
        logger.error(f"ocr_pdf failed: {e}")
        raise TextExtractionError("OCR extraction failed") from e
    finally:
        if pdf:
            pdf.close()
            
    if not text.strip():
        raise TextExtractionError("OCR extraction failed")
    return text.strip()


def extract_text(file_path: str, file_type: str) -> str:
    """Main dispatcher — extract text based on file type."""
    file_type = file_type.lower()
    logger.info(f"Starting text extraction for file type: {file_type}")
    
    text = ""
    try:
        if file_type in ['pdf', 'application/pdf']:
            text = extract_text_from_pdf(file_path)
        elif file_type in ['doc', 'docx', 'application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document']:
            if file_type in ['doc', 'application/msword'] or file_path.endswith('.doc'):
                text = extract_text_from_doc(file_path)
            else:
                text = extract_text_from_docx(file_path)
        elif file_type in ['txt', 'text/plain']:
            text = extract_text_from_txt(file_path)
        elif file_type in ['jpg', 'jpeg', 'png', 'bmp', 'tiff', 'webp', 
                          'image/jpeg', 'image/png', 'image/jpg', 'image/bmp', 'image/tiff', 'image/webp']:
            text = extract_text_from_image(file_path)
        else:
            raise TextExtractionError("Unsupported file structure")
    except TextExtractionError:
        raise
    except Exception as e:
        logger.error(f"Extraction failed: {e}", exc_info=True)
        raise TextExtractionError(f"Extraction failed: {e}") from e

    # OCR Fallback if text is empty
    if not text.strip():
        logger.info("Extracted text empty. Starting OCR fallback.")
        if file_type in ['pdf', 'application/pdf']:
            try:
                text = ocr_pdf(file_path)
            except TextExtractionError:
                raise
            except Exception as ocr_err:
                raise TextExtractionError(f"OCR extraction failed: {ocr_err}") from ocr_err
        elif file_type in ['jpg', 'jpeg', 'png', 'bmp', 'tiff', 'webp', 
                          'image/jpeg', 'image/png', 'image/jpg', 'image/bmp', 'image/tiff', 'image/webp']:
            # Do NOT run image OCR twice. If it returned empty in step 1, raise immediately
            raise TextExtractionError("No readable text found in image.")
        else:
            raise TextExtractionError("Unable to extract readable text from document.")
            
    if not text.strip():
        raise TextExtractionError("Unable to extract readable text from document.")
            
    logger.info(f"Text extraction successful. Character count: {len(text)}")
    return text.strip()


# ---------------------------------------------------------------------------
# Storage Usage Calculation (Unchanged Business Logic)
# ---------------------------------------------------------------------------
def get_user_storage_usage_mb(user_id: str) -> float:
    """Calculate the total storage used by a user in MB."""
    total_size_bytes = 0
    doc_ids_and_sizes = {}  # doc_id -> size_in_bytes

    # 1. Fetch documents from PostgreSQL
    try:
        from app.services.firebase_service import firebase_service
        conn = firebase_service._get_pg_conn()
        if conn:
            try:
                cur = conn.cursor()
                cur.execute("SELECT id, size_in_mb FROM documents WHERE user_id = %s", (user_id,))
                rows = cur.fetchall()
                for row in rows:
                    doc_id = row[0]
                    size_mb = row[1] or 0.0
                    doc_ids_and_sizes[doc_id] = int(size_mb * 1024 * 1024)
                cur.close()
                conn.close()
            except Exception as e:
                logger.error(f"Error querying PostgreSQL documents: {e}")
                if conn:
                    conn.close()
    except Exception as e:
        logger.error(f"PostgreSQL connection error: {e}")

    # 2. Fetch documents from Firestore
    try:
        from app.services.firebase_service import firebase_service
        if firebase_service.db:
            docs = firebase_service.db.collection("documents").where("user_id", "==", user_id).stream()
            for doc in docs:
                doc_id = doc.id
                data = doc.to_dict()
                size_mb = data.get("size_in_mb", 0.0) or 0.0
                if doc_id not in doc_ids_and_sizes:
                    doc_ids_and_sizes[doc_id] = int(size_mb * 1024 * 1024)
    except Exception as e:
        logger.error(f"Error querying Firestore documents: {e}")

    # 3. Fetch documents from Supabase Storage
    try:
        supabase = get_supabase()
        files = supabase.storage.from_("legal-documents").list(f"users/{user_id}/documents")
        if files:
            for f in files:
                name = f.get("name")
                if name:
                    doc_id, ext = os.path.splitext(name)
                    metadata = f.get('metadata')
                    if metadata:
                        size_bytes = metadata.get('size', 0)
                        doc_ids_and_sizes[doc_id] = size_bytes
    except Exception as e:
        logger.error(f"Error listing Supabase Storage files: {e}")

    # 4. Check actual file sizes on local disk
    from app.core.config import settings
    upload_dir = settings.UPLOAD_DIR
    local_doc_sizes = {}
    if os.path.exists(upload_dir):
        try:
            for filename in os.listdir(upload_dir):
                filepath = os.path.join(upload_dir, filename)
                if os.path.isfile(filepath):
                    doc_id, ext = os.path.splitext(filename)
                    if doc_id in doc_ids_and_sizes:
                        local_doc_sizes[doc_id] = os.path.getsize(filepath)
        except Exception as e:
            logger.error(f"Error scanning local upload directory: {e}")

    # For each found document ID, use local size if it exists, else the db/remote size
    for doc_id, remote_size in doc_ids_and_sizes.items():
        if doc_id in local_doc_sizes:
            total_size_bytes += local_doc_sizes[doc_id]
        else:
            total_size_bytes += remote_size

    # 5. Check generated reports
    reports_dir = os.path.join(upload_dir, "reports")
    if os.path.exists(reports_dir):
        try:
            for filename in os.listdir(reports_dir):
                filepath = os.path.join(reports_dir, filename)
                if os.path.isfile(filepath):
                    for doc_id in doc_ids_and_sizes.keys():
                        if f"LexGuard_Analysis_{doc_id}" in filename:
                            total_size_bytes += os.path.getsize(filepath)
                            break
        except Exception as e:
            logger.error(f"Error scanning local reports directory: {e}")

    total_size_mb = total_size_bytes / (1024 * 1024)
    return round(total_size_mb, 2)
